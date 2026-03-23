from __future__ import annotations
from typing import Dict, Any, Optional

def to_docx(result: Dict[str, Any], filepath: str) -> Optional[str]:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return None
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    heading = doc.add_heading('Medical Summary', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(result["summary_text"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('Chronological Timeline', level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Type"
    hdr[2].text = "Title"
    hdr[3].text = "Diagnoses"
    hdr[4].text = "Restrictions"

    for e in result["timeline"]:
        row = table.add_row().cells
        dt = e.get("date") or "Undated"
        row[0].text = str(dt)
        row[1].text = e["doc_type"]
        row[2].text = e["title"]
        row[3].text = ", ".join(e.get("diagnoses") or [])
        row[4].text = ", ".join(e.get("restrictions") or [])

    doc.save(filepath)
    return filepath